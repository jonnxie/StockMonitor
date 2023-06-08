import requests
import multitasking
import pandas as pd
from typing import List
from typing import Dict
import efinance as ef
from dataclasses import dataclass
from datetime import datetime
import rich
import docx

requests.packages.urllib3.util.ssl_.DEFAULT_CIPHERS = 'ALL:@SECLEVEL=1'


@dataclass()
class StockQuoteInfo:
    stock_code: str
    stock_name: str
    dt: datetime
    price: float
    top_price: float
    bottom_price: float
    latest_zt_dt: datetime
    latest_nzt_dt: datetime

    @property
    def zt_keep_seconds(self) -> int:
        return (self.latest_zt_dt - self.latest_nzt_dt).seconds


class Clock:
    def __init__(self) -> None:
        self.dt = datetime.now()

    def next(self) -> bool:
        dt = datetime.now()
        st = '09:15:00'
        et = '15:00:00'
        self.dt = dt
        return st <= dt.strftime('%H:%M:%S') <= et


def get_snapshot_fast(stock_codes: List[str]) -> Dict[str, pd.DataFrame]:
    sns: Dict[str, pd.DataFrame] = {}

    @multitasking.task
    def start(stock_code: str) -> None:
        sns[stock_code] = ef.stock.get_quote_snapshot(stock_code)

    for stock_code in stock_codes:
        start(stock_code)
    multitasking.wait_for_tasks()
    return sns


@dataclass()
class Monitor:
    stock_pool: Dict[int, str]
    stock_id_pool: List[int]

    def __init__(self) -> None:
        doc = docx.Document('热门股票.docx')
        self.stock_pool = doc.paragraphs


@dataclass()
class Strategy:
    clock: Clock

    def __post_init__(self) -> None:
        self.stock_code_info: Dict[str, StockQuoteInfo] = {}

    def next(self) -> None:
        dt = self.clock.dt

        quotes = ef.stock.get_realtime_quotes()
        quotes.index = quotes['股票代码'].values
        quotes = quotes[quotes['涨跌幅'] != '-']

        quotes = quotes[quotes['涨跌幅'] > 7]
        if len(quotes) == 0:
            return
        sns = get_snapshot_fast(quotes.index.values)
        for row in quotes.iloc:
            stock_code = row['股票代码']
            stock_name = row['股票名称']
            sn = sns[stock_code]
            top_price = sn['涨停价']
            bottom_price = sn['跌停价']
            current_price = sn['最新价']
            pre_info = self.stock_code_info.get(stock_code)
            first = pre_info is None
            if first:
                pre_info = StockQuoteInfo(
                    stock_code=stock_code,
                    stock_name=stock_name,
                    dt=dt,
                    price=current_price,
                    top_price=top_price,
                    bottom_price=bottom_price,
                    latest_nzt_dt=dt,
                    latest_zt_dt=None)
                self.stock_code_info[stock_code] = pre_info
            buy_list = []
            for i in range(1, 6):
                buy_list.append(f'买 {i}: {sn[f"买{i}数量"]}')
            buy_str = '\n'.join(buy_list)
            tip: str = None
            if abs(top_price - current_price) < 1e-2:
                if first or current_price > pre_info.price:
                    tip = ZT_TIP
                    pre_info.latest_zt_dt = dt
                elif current_price == pre_info.price:
                    tip = ZT_KEEP_TIP
                    pre_info.latest_zt_dt = dt
                else:
                    tip = ZT_BREAK_TIP
                    pre_info.latest_nzt_dt = dt
            else:
                pre_info.latest_nzt_dt = dt

            pre_info.price = current_price
            pre_info.dt = dt
            if tip == ZT_TIP or (tip == ZT_KEEP_TIP and pre_info.zt_keep_seconds <= ZT_NOTICE_MAX_SECONDS):
                msg = f'股票代码: {stock_code}\n股票名称: {stock_name}\n🚀 封单情况 🚀\n{buy_str}\n🚀 {tip} 🚀\n🚀 涨停保持秒数: {pre_info.zt_keep_seconds} 🚀'
                rich.print(msg)


TEST_MODE = True

ZT_TIP = '刚涨停'
ZT_KEEP_TIP = '保持涨停'
ZT_BREAK_TIP = '涨停炸板'
ZT_NOTICE_MAX_SECONDS = 60

monitor = Monitor
doc = docx.Document('热门股票.docx')
stock_pool: Dict[int, str] = doc.paragraphs

clock = Clock()
strategy = Strategy(clock)
while clock.next() or TEST_MODE:
    dt = clock.dt
    rich.print(f'[{dt.strftime("%m-%d %H:%M:%S")}] 刷新')
    strategy.next()
print('今日监控结束')
